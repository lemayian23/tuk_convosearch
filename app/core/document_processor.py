import os
from pathlib import Path
from typing import List, Dict, Any
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from loguru import logger
import PyPDF2
from docx import Document as DocxDocument
import json

class DocumentProcessor:
    def __init__(self, config):
        self.config = config
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.CHUNK_SIZE,
            chunk_overlap=config.CHUNK_OVERLAP,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def load_document(self, file_path: str) -> str:
        """Load document based on file type"""
        ext = Path(file_path).suffix.lower()
        
        try:
            if ext == '.pdf':
                return self._load_pdf(file_path)
            elif ext == '.docx':
                return self._load_docx(file_path)
            elif ext in ['.txt', '.md']:
                with open(file_path, 'r', encoding='utf-8') as f:
                    return f.read()
            elif ext == '.json':
                with open(file_path, 'r', encoding='utf-8') as f:
                    data = json.load(f)
                    return json.dumps(data, indent=2)
            else:
                logger.warning(f"Unsupported file type: {ext}")
                return ""
        except Exception as e:
            logger.error(f"Error loading {file_path}: {e}")
            return ""
    
    def _load_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def _load_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        doc = DocxDocument(file_path)
        return "\n".join([para.text for para in doc.paragraphs])
    
    def chunk_document(self, text: str, metadata: Dict) -> List[Document]:
        """Split document into chunks"""
        if not text.strip():
            return []
        
        # Create initial document
        doc = Document(page_content=text, metadata=metadata)
        
        # Split into chunks
        chunks = self.text_splitter.split_documents([doc])
        
        # Add chunk-specific metadata
        for i, chunk in enumerate(chunks):
            chunk.metadata.update({
                "chunk_id": i,
                "total_chunks": len(chunks)
            })
        
        logger.info(f"Created {len(chunks)} chunks from {metadata.get('source', 'unknown')}")
        return chunks
    
    def process_directory(self, directory_path: str) -> List[Document]:
        """Process all documents in a directory"""
        all_chunks = []
        directory = Path(directory_path)
        
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in ['.pdf', '.docx', '.txt', '.md', '.json']:
                logger.info(f"Processing: {file_path.name}")
                
                # Load document
                text = self.load_document(str(file_path))
                
                if text:
                    # Create metadata
                    metadata = {
                        "source": file_path.name,
                        "file_path": str(file_path),
                        "file_type": file_path.suffix.lower(),
                        "file_size": file_path.stat().st_size
                    }
                    
                    # Chunk document
                    chunks = self.chunk_document(text, metadata)
                    all_chunks.extend(chunks)
        
        logger.info(f"Total chunks created: {len(all_chunks)}")
        return all_chunks